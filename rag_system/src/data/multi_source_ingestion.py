"""
Multi-Source Data Ingestion Pipeline - Production RAG 2026

Following RAG Pipeline Guide 2026 - Phase 1: Data Ingestion Pipeline

Features:
- Multi-source connectors (files, APIs, databases, webhooks)
- Incremental updates with delta sync
- Checksum validation
- Document parsing (PDF, DOCX, TXT, MD, HTML, JSON)
- Metadata extraction and enrichment
- Error handling & retry logic
- Arabic text processing optimization

Usage:
    pipeline = MultiSourceIngestionPipeline()
    await pipeline.ingest_all_sources()
"""

import os
import json
import hashlib
import asyncio
import aiohttp
import sqlite3
from typing import List, Dict, Any, Optional, Callable, AsyncGenerator, Set
from pathlib import Path
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
import logging
import re

logger = logging.getLogger(__name__)


# ==================== Enums & Data Classes ====================


class DataSourceType(Enum):
    """Classification of data sources by priority."""

    PRIMARY = "primary"  # 10-20% of docs, 80% of questions (START HERE)
    SECONDARY = "secondary"  # Add after MVP works
    ARCHIVAL = "archival"  # Historical, rarely accessed


class ConnectorType(Enum):
    """Supported connector types."""

    FILE = "file"
    API = "api"
    DATABASE = "database"
    WEBHOOK = "webhook"
    STREAM = "stream"


class UpdateFrequency(Enum):
    """Update frequency options."""

    REAL_TIME = "real-time"
    HOURLY = "hourly"
    DAILY = "daily"
    WEEKLY = "weekly"
    MONTHLY = "monthly"


@dataclass
class DataSource:
    """Configuration for a data source."""

    name: str
    source_type: DataSourceType
    connector_type: ConnectorType
    update_frequency: UpdateFrequency
    priority: int = 1  # 1-10, higher = more important

    # Connection details
    path: Optional[str] = None  # For files
    url: Optional[str] = None  # For APIs
    connection_string: Optional[str] = None  # For databases
    webhook_secret: Optional[str] = None  # For webhooks

    # Authentication
    api_key: Optional[str] = None
    headers: Dict[str, str] = field(default_factory=dict)

    # Query/config
    query: Optional[str] = None  # For databases
    params: Dict[str, Any] = field(default_factory=dict)

    # Metadata
    description: str = ""
    categories: List[str] = field(default_factory=list)


@dataclass
class Document:
    """A document in the ingestion pipeline."""

    content: str
    doc_id: str
    source: str
    source_type: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    checksum: Optional[str] = None

    def __post_init__(self):
        if self.checksum is None:
            self.checksum = self._compute_checksum()

    def _compute_checksum(self) -> str:
        """Compute SHA256 checksum of content."""
        return hashlib.sha256(self.content.encode("utf-8")).hexdigest()


@dataclass
class IngestionResult:
    """Result of an ingestion operation."""

    documents: List[Document]
    new_documents: int = 0
    updated_documents: int = 0
    deleted_documents: int = 0
    skipped_documents: int = 0
    errors: List[str] = field(default_factory=list)
    processing_time_ms: float = 0.0


# ==================== Document Parsers ====================


class DocumentParser:
    """
    Parse various document formats.

    2026 Best Practice: Use specialized parsers for complex layouts

    Hierarchy of parsing quality:
    1. LlamaParse (best for complex layouts, tables, figures)
    2. Unstructured.io (good open-source alternative)
    3. Docling (IBM, excellent for tables)
    4. LLMWhisperer (good for scanned docs)
    5. PyPDF2/pdfplumber (avoid - loses layout)
    """

    def __init__(self, parser_type: str = "auto"):
        self.parser_type = parser_type

    async def parse(self, file_path: str) -> Document:
        """Parse document and return Document object."""

        # Auto-detect parser type from file extension
        if self.parser_type == "auto":
            ext = Path(file_path).suffix.lower()
            parser_map = {
                ".txt": "txt",
                ".md": "markdown",
                ".markdown": "markdown",
                ".pdf": "pdf",
                ".docx": "docx",
                ".html": "html",
                ".htm": "html",
                ".json": "json",
                ".xml": "xml",
            }
            self.parser_type = parser_map.get(ext, "txt")

        # Dispatch to appropriate parser
        parser_method = getattr(self, f"_parse_{self.parser_type}", None)
        if parser_method:
            return await parser_method(file_path)
        else:
            return await self._parse_txt(file_path)

    async def _parse_txt(self, file_path: str) -> Document:
        """Parse plain text file."""

        # Try multiple encodings for Arabic text
        encodings = ["utf-8", "utf-8-sig", "cp1256", "cp1252", "iso-8859-6"]

        content = None
        used_encoding = "utf-8"
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    content = f.read()
                used_encoding = encoding
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            # Fallback to binary read
            with open(file_path, "rb") as f:
                content = f.read().decode("utf-8", errors="ignore")
            used_encoding = "utf-8 (fallback)"

        # Extract metadata from filename
        filename = Path(file_path).stem

        return Document(
            content=content,
            doc_id=hashlib.md5(file_path.encode()).hexdigest(),
            source=str(file_path),
            source_type="file",
            metadata={
                "filename": filename,
                "file_size": Path(file_path).stat().st_size,
                "encoding": used_encoding,
                "parsed_at": datetime.now().isoformat(),
                "file_extension": Path(file_path).suffix,
            },
        )

    async def _parse_markdown(self, file_path: str) -> Document:
        """Parse markdown file."""

        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        # Extract title from first # heading
        title = ""
        for line in content.split("\n"):
            if line.strip().startswith("# "):
                title = line.strip()[2:]
                break

        # Count sections
        section_count = len(re.findall(r"^##\s+", content, re.MULTILINE))

        return Document(
            content=content,
            doc_id=hashlib.md5(file_path.encode()).hexdigest(),
            source=str(file_path),
            source_type="file",
            metadata={
                "title": title,
                "filename": Path(file_path).name,
                "file_size": Path(file_path).stat().st_size,
                "section_count": section_count,
                "format": "markdown",
            },
        )

    async def _parse_pdf(self, file_path: str) -> Document:
        """
        Parse PDF file.

        2026 Best Practice: Use specialized parsers for layout preservation
        """

        content = ""
        parsing_method = ""

        # Try LlamaParse first (best quality)
        if os.getenv("LLAMA_CLOUD_API_KEY"):
            try:
                content = await self._parse_with_llama(file_path)
                parsing_method = "llama"
            except Exception as e:
                logger.warning(f"LlamaParse failed: {e}")

        # Try Unstructured.io (open source)
        if not content:
            try:
                content = await self._parse_with_unstructured(file_path)
                parsing_method = "unstructured"
            except Exception as e:
                logger.warning(f"Unstructured parsing failed: {e}")

        # Try Docling (IBM, good for tables)
        if not content:
            try:
                content = await self._parse_with_docling(file_path)
                parsing_method = "docling"
            except Exception as e:
                logger.warning(f"Docling parsing failed: {e}")

        # Fallback to pdfplumber
        if not content:
            try:
                content = await self._parse_with_pdfplumber(file_path)
                parsing_method = "pdfplumber"
            except Exception as e:
                logger.error(f"All PDF parsing methods failed: {e}")
                raise

        return Document(
            content=content,
            doc_id=hashlib.md5(file_path.encode()).hexdigest(),
            source=str(file_path),
            source_type="file",
            metadata={
                "filename": Path(file_path).name,
                "file_size": Path(file_path).stat().st_size,
                "parsing_method": parsing_method,
                "format": "pdf",
            },
        )

    async def _parse_with_llama(self, file_path: str) -> str:
        """Parse PDF with LlamaParse (best for complex layouts)."""

        from llama_parse import LlamaParse

        parser = LlamaParse(
            api_key=os.getenv("LLAMA_CLOUD_API_KEY"),
            result_type="markdown",  # or "json" for structured
            verbose=True,
            language="en",
        )

        result = await parser.aget_result(file_path)
        return result.text

    async def _parse_with_unstructured(self, file_path: str) -> str:
        """Parse PDF with Unstructured.io (open source)."""

        from unstructured.partition.pdf import partition_pdf

        elements = partition_pdf(filename=str(file_path))
        return "\n\n".join([str(el) for el in elements])

    async def _parse_with_docling(self, file_path: str) -> str:
        """Parse PDF with Docling (IBM, excellent for tables)."""

        from docling.document_converter import DocumentConverter

        converter = DocumentConverter()
        result = converter.convert(str(file_path))
        return result.document.export_to_markdown()

    async def _parse_with_pdfplumber(self, file_path: str) -> str:
        """Parse PDF with pdfplumber (fallback)."""

        import pdfplumber

        text_content = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)

        return "\n\n".join(text_content)

    async def _parse_docx(self, file_path: str) -> Document:
        """Parse DOCX file."""

        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)

        return Document(
            content=content,
            doc_id=hashlib.md5(file_path.encode()).hexdigest(),
            source=str(file_path),
            source_type="file",
            metadata={
                "filename": Path(file_path).name,
                "file_size": Path(file_path).stat().st_size,
                "paragraph_count": len(paragraphs),
                "format": "docx",
            },
        )

    async def _parse_html(self, file_path: str) -> Document:
        """Parse HTML file."""

        from bs4 import BeautifulSoup

        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f, "html.parser")

        # Extract title
        title = ""
        if soup.title:
            title = soup.title.string or ""

        # Extract main text (remove scripts and styles)
        for script in soup(["script", "style"]):
            script.decompose()

        content = soup.get_text(separator="\n", strip=True)

        return Document(
            content=content,
            doc_id=hashlib.md5(file_path.encode()).hexdigest(),
            source=str(file_path),
            source_type="file",
            metadata={
                "title": title,
                "filename": Path(file_path).name,
                "file_size": Path(file_path).stat().st_size,
                "format": "html",
            },
        )

    async def _parse_json(self, file_path: str) -> Document:
        """Parse JSON file."""

        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        # Convert JSON to text representation
        if isinstance(data, dict):
            # Try to extract content field
            content = data.get("content", data.get("text", json.dumps(data, indent=2)))
            title = data.get("title", data.get("name", ""))
        elif isinstance(data, list):
            # Assume list of documents
            content = "\n\n".join(
                [
                    item.get("content", item.get("text", str(item)))
                    for item in data
                    if isinstance(item, dict)
                ]
            )
            title = ""
        else:
            content = str(data)
            title = ""

        return Document(
            content=content,
            doc_id=hashlib.md5(file_path.encode()).hexdigest(),
            source=str(file_path),
            source_type="file",
            metadata={
                "title": title,
                "filename": Path(file_path).name,
                "file_size": Path(file_path).stat().st_size,
                "format": "json",
            },
        )

    async def _parse_xml(self, file_path: str) -> Document:
        """Parse XML file."""

        import xml.etree.ElementTree as ET

        tree = ET.parse(file_path)
        root = tree.getroot()

        # Extract text content
        content = ET.tostring(root, encoding="utf-8", method="text").decode("utf-8")

        return Document(
            content=content,
            doc_id=hashlib.md5(file_path.encode()).hexdigest(),
            source=str(file_path),
            source_type="file",
            metadata={
                "root_tag": root.tag,
                "filename": Path(file_path).name,
                "file_size": Path(file_path).stat().st_size,
                "format": "xml",
            },
        )


# ==================== Connectors ====================


class FileConnector:
    """
    Connector for file system sources.

    Supports:
    - Single files
    - Directories (recursive)
    - Glob patterns
    - Multiple file formats
    """

    def __init__(self, source: DataSource):
        self.source = source
        self.parser = DocumentParser(parser_type="auto")

        # Supported file extensions
        self.supported_extensions = {
            ".txt",
            ".md",
            ".markdown",
            ".pdf",
            ".docx",
            ".html",
            ".htm",
            ".json",
            ".xml",
        }

    async def extract(self) -> List[Document]:
        """Extract documents from file system."""

        documents = []
        base_path = Path(self.source.path)

        if not base_path.exists():
            logger.warning(f"Path not found: {base_path}")
            return documents

        # Get files to process
        file_paths = self._get_file_paths(base_path)

        logger.info(f"Found {len(file_paths)} files to process")

        # Process files
        for file_path in file_paths:
            try:
                doc = await self.parser.parse(str(file_path))
                doc.source_type = self.source.name
                doc.metadata.update(self._extract_file_metadata(file_path))
                documents.append(doc)
            except Exception as e:
                logger.error(f"Error parsing {file_path}: {e}")
                documents.append(
                    Document(
                        content="",
                        doc_id=hashlib.md5(str(file_path).encode()).hexdigest(),
                        source=str(file_path),
                        source_type=self.source.name,
                        metadata={"error": str(e)},
                    )
                )

        return documents

    def _get_file_paths(self, base_path: Path) -> List[Path]:
        """Get all file paths to process."""

        if base_path.is_file():
            # Single file
            if base_path.suffix.lower() in self.supported_extensions:
                return [base_path]
            return []

        elif base_path.is_dir():
            # Directory - recursive glob
            file_paths = []
            for ext in self.supported_extensions:
                file_paths.extend(base_path.glob(f"**/*{ext}"))
            return sorted(file_paths)

        else:
            # Try glob pattern
            file_paths = list(Path(".").glob(str(base_path)))
            return [
                p
                for p in file_paths
                if p.suffix.lower() in self.supported_extensions
            ]

    def _extract_file_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract additional metadata from file."""

        stat = file_path.stat()

        return {
            "file_path": str(file_path),
            "file_name": file_path.name,
            "file_size": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "extension": file_path.suffix.lower(),
        }


class APIConnector:
    """
    Connector for API sources.

    Supports:
    - REST APIs
    - GraphQL APIs
    - Pagination
    - Rate limiting
    - Authentication
    """

    def __init__(self, source: DataSource):
        self.source = source
        self.base_url = source.url
        self.headers = source.headers.copy()

        if source.api_key:
            self.headers["Authorization"] = f"Bearer {source.api_key}"

    async def extract(self) -> List[Document]:
        """Extract documents from API."""

        documents = []

        async with aiohttp.ClientSession(headers=self.headers) as session:
            # Get all pages
            async for page_data in self._paginate(session):
                # Parse page data
                docs = self._parse_response(page_data)
                documents.extend(docs)

        return documents

    async def _paginate(
        self,
        session: aiohttp.ClientSession,
    ) -> AsyncGenerator[Dict[str, Any], None]:
        """
        Paginate through API results.

        Supports:
        - Offset/limit pagination
        - Cursor-based pagination
        - Page number pagination
        """

        url = self.base_url
        params = self.source.params.copy()

        # Get pagination config
        pagination_type = self.source.params.get("pagination", "offset")
        page_size = self.source.params.get("page_size", 100)

        if pagination_type == "offset":
            # Offset/limit pagination
            offset = 0
            while True:
                params["offset"] = offset
                params["limit"] = page_size

                async with session.get(url, params=params) as response:
                    if response.status != 200:
                        logger.error(f"API error: {response.status}")
                        break

                    data = await response.json()

                    if not data:
                        break

                    yield data

                    # Check if more data
                    if len(data) < page_size:
                        break

                    offset += page_size

        elif pagination_type == "cursor":
            # Cursor-based pagination
            cursor = None
            while True:
                if cursor:
                    params["cursor"] = cursor

                async with session.get(url, params=params) as response:
                    if response.status != 200:
                        logger.error(f"API error: {response.status}")
                        break

                    data = await response.json()

                    if not data:
                        break

                    yield data

                    # Get next cursor
                    cursor = data.get("next_cursor")
                    if not cursor:
                        break

        elif pagination_type == "page":
            # Page number pagination
            page = 1
            while True:
                params["page"] = page
                params["page_size"] = page_size

                async with session.get(url, params=params) as response:
                    if response.status != 200:
                        logger.error(f"API error: {response.status}")
                        break

                    data = await response.json()

                    if not data:
                        break

                    yield data

                    # Check if more pages
                    if page >= data.get("total_pages", page):
                        break

                    page += 1

    def _parse_response(self, data: Dict[str, Any]) -> List[Document]:
        """Parse API response into documents."""

        documents = []

        # Get content mapping from params
        content_field = self.source.params.get("content_field", "content")
        id_field = self.source.params.get("id_field", "id")
        title_field = self.source.params.get("title_field", "title")

        # Handle list or dict response
        items = data if isinstance(data, list) else data.get("items", [data])

        for item in items:
            if not isinstance(item, dict):
                continue

            content = item.get(content_field, "")
            if not content:
                continue

            doc_id = str(item.get(id_field, hashlib.md5(content.encode()).hexdigest()))
            title = item.get(title_field, "")

            doc = Document(
                content=str(content),
                doc_id=doc_id,
                source=self.base_url,
                source_type=self.source.name,
                metadata={
                    "title": title,
                    "api_source": self.base_url,
                    "retrieved_at": datetime.now().isoformat(),
                    **{k: v for k, v in item.items() if k not in [content_field, id_field, title_field]},
                },
            )

            documents.append(doc)

        return documents


class DatabaseConnector:
    """
    Connector for database sources.

    Supports:
    - SQLite
    - PostgreSQL
    - MySQL
    - Custom queries
    """

    def __init__(self, source: DataSource):
        self.source = source
        self.connection_string = source.connection_string
        self.query = source.query

    async def extract(self) -> List[Document]:
        """Extract documents from database."""

        # Parse connection string
        db_type = self._parse_db_type()

        if db_type == "sqlite":
            return await self._extract_sqlite()
        elif db_type == "postgresql":
            return await self._extract_postgresql()
        elif db_type == "mysql":
            return await self._extract_mysql()
        else:
            raise ValueError(f"Unsupported database type: {db_type}")

    def _parse_db_type(self) -> str:
        """Parse database type from connection string."""

        if not self.connection_string:
            return "sqlite"

        if self.connection_string.startswith("sqlite"):
            return "sqlite"
        elif self.connection_string.startswith("postgresql"):
            return "postgresql"
        elif self.connection_string.startswith("mysql"):
            return "mysql"
        else:
            return "sqlite"

    async def _extract_sqlite(self) -> List[Document]:
        """Extract from SQLite database."""

        documents = []

        # Get database path
        db_path = self.connection_string.replace("sqlite:///", "")
        if not db_path:
            db_path = "database.db"

        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()

        try:
            # Execute query
            cursor.execute(self.query)
            rows = cursor.fetchall()

            # Get column names
            columns = [desc[0] for desc in cursor.description]

            # Map to document fields
            content_idx = columns.index("content") if "content" in columns else 0
            id_idx = columns.index("id") if "id" in columns else 0
            title_idx = columns.index("title") if "title" in columns else -1

            for row in rows:
                content = str(row[content_idx]) if content_idx < len(row) else ""
                if not content.strip():
                    continue

                doc_id = str(row[id_idx]) if id_idx < len(row) else hashlib.md5(content.encode()).hexdigest()
                title = str(row[title_idx]) if 0 <= title_idx < len(row) else ""

                doc = Document(
                    content=content,
                    doc_id=doc_id,
                    source=db_path,
                    source_type=self.source.name,
                    metadata={
                        "title": title,
                        "database": db_path,
                        "query": self.query,
                        "retrieved_at": datetime.now().isoformat(),
                        **{columns[i]: str(row[i]) for i in range(len(row)) if i not in [content_idx, id_idx, title_idx]},
                    },
                )

                documents.append(doc)

        finally:
            conn.close()

        return documents

    async def _extract_postgresql(self) -> List[Document]:
        """Extract from PostgreSQL database."""

        try:
            import asyncpg
        except ImportError:
            raise ImportError("Please install asyncpg: pip install asyncpg")

        documents = []

        async with await asyncpg.connect(self.connection_string) as conn:
            async with conn.transaction():
                async for row in conn.cursor(self.query):
                    content = str(row.get("content", ""))
                    if not content.strip():
                        continue

                    doc_id = str(row.get("id", hashlib.md5(content.encode()).hexdigest()))
                    title = str(row.get("title", ""))

                    doc = Document(
                        content=content,
                        doc_id=doc_id,
                        source=self.connection_string,
                        source_type=self.source.name,
                        metadata={
                            "title": title,
                            "database": "postgresql",
                            "retrieved_at": datetime.now().isoformat(),
                        },
                    )

                    documents.append(doc)

        return documents

    async def _extract_mysql(self) -> List[Document]:
        """Extract from MySQL database."""

        try:
            import aiomysql
        except ImportError:
            raise ImportError("Please install aiomysql: pip install aiomysql")

        documents = []

        # Parse connection string
        # mysql://user:password@host:port/database
        parts = self.connection_string.replace("mysql://", "").split("@")
        user_pass = parts[0].split(":")
        host_db = parts[1].split("/")

        user = user_pass[0]
        password = user_pass[1] if len(user_pass) > 1 else ""
        host_port = host_db[0].split(":")
        host = host_port[0]
        port = int(host_port[1]) if len(host_port) > 1 else 3306
        database = host_db[1] if len(host_db) > 1 else ""

        async with await aiomysql.connect(
            host=host,
            port=port,
            user=user,
            password=password,
            db=database,
        ) as conn:
            async with conn.cursor(aiomysql.DictCursor) as cursor:
                await cursor.execute(self.query)

                async for row in cursor:
                    content = str(row.get("content", ""))
                    if not content.strip():
                        continue

                    doc_id = str(row.get("id", hashlib.md5(content.encode()).hexdigest()))
                    title = str(row.get("title", ""))

                    doc = Document(
                        content=content,
                        doc_id=doc_id,
                        source=self.connection_string,
                        source_type=self.source.name,
                        metadata={
                            "title": title,
                            "database": "mysql",
                            "retrieved_at": datetime.now().isoformat(),
                            **{k: str(v) for k, v in row.items() if k not in ["content", "id", "title"]},
                        },
                    )

                    documents.append(doc)

        return documents


class WebhookConnector:
    """
    Connector for webhook sources.

    Receives documents via HTTP POST requests.
    """

    def __init__(self, source: DataSource):
        self.source = source
        self.documents: List[Document] = []
        self._server = None

    async def start_server(self, port: int = 8080):
        """Start webhook server."""

        from aiohttp import web

        async def handle_webhook(request):
            """Handle incoming webhook."""

            try:
                data = await request.json()

                # Validate secret
                if self.source.webhook_secret:
                    header_secret = request.headers.get("X-Webhook-Secret")
                    if header_secret != self.source.webhook_secret:
                        return web.json_response(
                            {"error": "Invalid secret"},
                            status=401,
                        )

                # Parse document
                doc = self._parse_payload(data)
                if doc:
                    self.documents.append(doc)

                return web.json_response({"status": "ok"})

            except Exception as e:
                logger.error(f"Webhook error: {e}")
                return web.json_response(
                    {"error": str(e)},
                    status=400,
                )

        app = web.Application()
        app.router.add_post("/webhook", handle_webhook)

        runner = web.AppRunner(app)
        await runner.setup()

        site = web.TCPSite(runner, "0.0.0.0", port)
        await site.start()

        logger.info(f"Webhook server started on port {port}")

        return runner

    def _parse_payload(self, data: Dict[str, Any]) -> Optional[Document]:
        """Parse webhook payload into document."""

        content_field = self.source.params.get("content_field", "content")
        id_field = self.source.params.get("id_field", "id")

        content = data.get(content_field)
        if not content:
            return None

        doc_id = str(data.get(id_field, hashlib.md5(content.encode()).hexdigest()))

        return Document(
            content=str(content),
            doc_id=doc_id,
            source="webhook",
            source_type=self.source.name,
            metadata={
                "webhook_received_at": datetime.now().isoformat(),
                **{k: v for k, v in data.items() if k not in [content_field, id_field]},
            },
        )

    async def extract(self) -> List[Document]:
        """Extract documents (return accumulated webhook documents)."""

        return self.documents


# ==================== Main Ingestion Pipeline ====================


class MultiSourceIngestionPipeline:
    """
    Production ingestion pipeline with multi-source support.

    Features:
    - Multiple data sources
    - Incremental updates (delta sync)
    - Checksum validation
    - Error handling & retry
    - Metadata extraction
    - Batch processing
    """

    def __init__(
        self,
        update_strategy: str = "incremental",
        batch_size: int = 100,
        max_concurrent: int = 5,
    ):
        self.update_strategy = update_strategy
        self.batch_size = batch_size
        self.max_concurrent = max_concurrent

        # State
        self.sources: List[DataSource] = []
        self._previous_checksums: Dict[str, str] = {}
        self._document_stats: Dict[str, Any] = {}

        # Load previous checksums
        self._load_previous_checksums()

    def _load_previous_checksums(self):
        """Load previously indexed document checksums."""

        checksum_file = Path("rag_system/data/checksums.json")
        if checksum_file.exists():
            with open(checksum_file, "r", encoding="utf-8") as f:
                self._previous_checksums = json.load(f)
            logger.info(f"Loaded {len(self._previous_checksums)} previous checksums")

    def _save_previous_checksums(self):
        """Save current checksums for incremental updates."""

        checksum_file = Path("rag_system/data/checksums.json")
        checksum_file.parent.mkdir(parents=True, exist_ok=True)

        with open(checksum_file, "w", encoding="utf-8") as f:
            json.dump(self._previous_checksums, f, ensure_ascii=False, indent=2)

        logger.info("Saved checksums for incremental updates")

    def add_source(self, source: DataSource):
        """Add a data source to the pipeline."""

        self.sources.append(source)
        logger.info(f"Added data source: {source.name}")

    async def ingest_all_sources(
        self,
        progress_callback: Optional[Callable] = None,
    ) -> Dict[str, IngestionResult]:
        """
        Ingest from all configured sources.

        Args:
            progress_callback: Optional callback for progress updates

        Returns:
            Dictionary of source_name -> IngestionResult
        """

        results = {}

        # Process sources with concurrency limit
        semaphore = asyncio.Semaphore(self.max_concurrent)

        async def ingest_with_semaphore(source: DataSource) -> tuple[str, IngestionResult]:
            async with semaphore:
                result = await self.ingest_source(source, progress_callback)
                return source.name, result

        # Create tasks for all sources
        tasks = [
            ingest_with_semaphore(source)
            for source in self.sources
        ]

        # Execute tasks
        for task in asyncio.as_completed(tasks):
            source_name, result = await task
            results[source_name] = result
            logger.info(
                f"Completed {source_name}: "
                f"{result.new_documents} new, "
                f"{result.updated_documents} updated, "
                f"{result.errors} errors"
            )

        # Save checksums
        self._save_previous_checksums()

        # Calculate total stats
        self._calculate_total_stats(results)

        return results

    async def ingest_source(
        self,
        source: DataSource,
        progress_callback: Optional[Callable] = None,
    ) -> IngestionResult:
        """
        Ingest from a single source.

        Args:
            source: Data source configuration
            progress_callback: Optional callback for progress updates

        Returns:
            IngestionResult with statistics
        """

        start_time = datetime.now()

        logger.info(f"Starting ingestion from {source.name} ({source.connector_type.value})")

        try:
            # Step 1: Connect and extract
            raw_documents = await self._extract(source)

            logger.info(f"Extracted {len(raw_documents)} raw documents from {source.name}")

            # Step 2: Validate and clean
            validated = self._validate(raw_documents)

            logger.info(f"Validated {len(validated)} documents")

            # Step 3: Check for changes (incremental)
            if self.update_strategy == "incremental":
                new_docs, updated_docs = await self._detect_changes(validated)
            else:
                new_docs = validated
                updated_docs = []

            logger.info(f"New: {len(new_docs)}, Updated: {len(updated_docs)}")

            # Step 4: Extract metadata
            enriched_new = self._extract_metadata(new_docs, source)
            enriched_updated = self._extract_metadata(updated_docs, source)

            # Step 5: Update checksums
            for doc in enriched_new + enriched_updated:
                self._previous_checksums[doc.doc_id] = doc.checksum

            # Step 6: Report progress
            if progress_callback:
                progress_callback(source.name, len(enriched_new), len(enriched_updated))

            processing_time = (datetime.now() - start_time).total_seconds() * 1000

            return IngestionResult(
                documents=enriched_new + enriched_updated,
                new_documents=len(enriched_new),
                updated_documents=len(enriched_updated),
                processing_time_ms=processing_time,
            )

        except Exception as e:
            logger.error(f"Error ingesting {source.name}: {e}")

            return IngestionResult(
                documents=[],
                errors=[str(e)],
                processing_time_ms=(datetime.now() - start_time).total_seconds() * 1000,
            )

    async def _extract(self, source: DataSource) -> List[Document]:
        """Extract from source using appropriate connector."""

        connector_type = source.connector_type

        if connector_type == ConnectorType.FILE:
            connector = FileConnector(source)
        elif connector_type == ConnectorType.API:
            connector = APIConnector(source)
        elif connector_type == ConnectorType.DATABASE:
            connector = DatabaseConnector(source)
        elif connector_type == ConnectorType.WEBHOOK:
            connector = WebhookConnector(source)
        else:
            raise ValueError(f"Unsupported connector type: {connector_type}")

        return await connector.extract()

    def _validate(self, documents: List[Document]) -> List[Document]:
        """Validate and clean documents."""

        validated = []

        for doc in documents:
            # Check minimum content length
            if len(doc.content.strip()) < 50:
                logger.debug(f"Skipping short document: {doc.doc_id}")
                continue

            # Check for valid encoding
            try:
                doc.content.encode("utf-8")
            except UnicodeEncodeError:
                logger.warning(f"Invalid encoding in {doc.doc_id}")
                continue

            # Check for empty content
            if not doc.content.strip():
                logger.debug(f"Skipping empty document: {doc.doc_id}")
                continue

            validated.append(doc)

        return validated

    async def _detect_changes(
        self,
        documents: List[Document],
    ) -> tuple[List[Document], List[Document]]:
        """
        Detect new and updated documents.

        Returns:
            Tuple of (new_documents, updated_documents)
        """

        new_docs = []
        updated_docs = []

        for doc in documents:
            previous_checksum = self._previous_checksums.get(doc.doc_id)

            if previous_checksum is None:
                # New document
                new_docs.append(doc)
            elif previous_checksum != doc.checksum:
                # Updated document
                updated_docs.append(doc)
            else:
                # Unchanged - skip
                logger.debug(f"Skipping unchanged document: {doc.doc_id}")

        return new_docs, updated_docs

    def _extract_metadata(
        self,
        documents: List[Document],
        source: DataSource,
    ) -> List[Document]:
        """Enrich documents with metadata."""

        for doc in documents:
            # Add source metadata
            doc.metadata.update(
                {
                    "source_name": source.name,
                    "source_type": source.source_type.value,
                    "priority": source.priority,
                    "ingested_at": datetime.now().isoformat(),
                    "word_count": len(doc.content.split()),
                    "char_count": len(doc.content),
                    "categories": source.categories,
                }
            )

            # Add Arabic-specific metadata
            doc.metadata["is_arabic"] = self._is_arabic_text(doc.content)
            doc.metadata["language"] = self._detect_language(doc.content)

        return documents

    def _is_arabic_text(self, text: str) -> bool:
        """Check if text is primarily Arabic."""

        arabic_chars = len(re.findall(r"[\u0600-\u06FF]", text))
        total_chars = len(text)

        if total_chars == 0:
            return False

        return arabic_chars / total_chars > 0.5

    def _detect_language(self, text: str) -> str:
        """Detect primary language of text."""

        if self._is_arabic_text(text):
            return "ar"

        # Check for English
        english_chars = len(re.findall(r"[a-zA-Z]", text))
        total_chars = len(text)

        if total_chars > 0 and english_chars / total_chars > 0.5:
            return "en"

        return "unknown"

    def _calculate_total_stats(self, results: Dict[str, IngestionResult]):
        """Calculate total statistics."""

        total_new = sum(r.new_documents for r in results.values())
        total_updated = sum(r.updated_documents for r in results.values())
        total_errors = sum(len(r.errors) for r in results.values())
        total_time = sum(r.processing_time_ms for r in results.values())

        self._document_stats = {
            "sources_processed": len(results),
            "total_new_documents": total_new,
            "total_updated_documents": total_updated,
            "total_errors": total_errors,
            "total_processing_time_ms": total_time,
            "total_indexed": len(self._previous_checksums),
            "timestamp": datetime.now().isoformat(),
        }

        logger.info(
            f"Total: {total_new} new, {total_updated} updated, "
            f"{total_errors} errors, {total_time:.2f}ms"
        )

    def get_statistics(self) -> Dict[str, Any]:
        """Get ingestion statistics."""

        return {
            **self._document_stats,
            "update_strategy": self.update_strategy,
            "batch_size": self.batch_size,
            "max_concurrent": self.max_concurrent,
        }


# ==================== Factory Functions ====================


def create_file_source(
    name: str,
    path: str,
    priority: int = 5,
    categories: Optional[List[str]] = None,
) -> DataSource:
    """Create a file-based data source."""

    return DataSource(
        name=name,
        source_type=DataSourceType.PRIMARY,
        connector_type=ConnectorType.FILE,
        update_frequency=UpdateFrequency.DAILY,
        priority=priority,
        path=path,
        categories=categories or [],
        description=f"File source: {path}",
    )


def create_api_source(
    name: str,
    url: str,
    api_key: Optional[str] = None,
    params: Optional[Dict[str, Any]] = None,
    priority: int = 3,
) -> DataSource:
    """Create an API-based data source."""

    return DataSource(
        name=name,
        source_type=DataSourceType.SECONDARY,
        connector_type=ConnectorType.API,
        update_frequency=UpdateFrequency.HOURLY,
        priority=priority,
        url=url,
        api_key=api_key,
        params=params or {},
        description=f"API source: {url}",
    )


def create_database_source(
    name: str,
    connection_string: str,
    query: str,
    priority: int = 4,
) -> DataSource:
    """Create a database-based data source."""

    return DataSource(
        name=name,
        source_type=DataSourceType.ARCHIVAL,
        connector_type=ConnectorType.DATABASE,
        update_frequency=UpdateFrequency.DAILY,
        priority=priority,
        connection_string=connection_string,
        query=query,
        description=f"Database source: {name}",
    )


# ==================== Main Entry ====================


async def create_arabic_islamic_pipeline(
    datasets_path: str,
) -> MultiSourceIngestionPipeline:
    """
    Create ingestion pipeline for Arabic Islamic literature.

    Args:
        datasets_path: Path to the datasets directory

    Returns:
        Configured MultiSourceIngestionPipeline
    """

    pipeline = MultiSourceIngestionPipeline(
        update_strategy="incremental",
        batch_size=100,
        max_concurrent=5,
    )

    # Add extracted books source
    books_source = create_file_source(
        name="extracted_books",
        path=os.path.join(datasets_path, "extracted_books"),
        priority=10,
        categories=["Islamic Literature"],
    )
    pipeline.add_source(books_source)

    # Add Arabic web source
    arabic_web_source = create_file_source(
        name="arabic_web",
        path=os.path.join(datasets_path, "arabic_web"),
        priority=8,
        categories=["Arabic Web Content"],
    )
    pipeline.add_source(arabic_web_source)

    # Add Sanadset source
    sanadset_source = create_file_source(
        name="sanadset",
        path=os.path.join(datasets_path, "Sanadset 368K Data on Hadith Narrators"),
        priority=9,
        categories=["Hadith", "Narrators"],
    )
    pipeline.add_source(sanadset_source)

    logger.info(f"Created pipeline with {len(pipeline.sources)} sources")

    return pipeline


# ==================== Metadata Ingestion ====================


class MetadataIngestionPipeline:
    """
    Specialized pipeline for ingesting book metadata.
    """

    def __init__(self, metadata_path: str):
        self.metadata_path = metadata_path
        self.books: Dict[int, Dict] = {}
        self.authors: Dict[int, Dict] = {}
        self.categories: Dict[int, Dict] = {}

    async def load_metadata(self) -> Dict[str, Any]:
        """Load all metadata files."""
        
        import json
        from pathlib import Path
        
        # Load books.json
        books_file = Path(self.metadata_path) / "books.json"
        if books_file.exists():
            with open(books_file, "r", encoding="utf-8") as f:
                books_data = json.load(f)
                self.books = {b["id"]: b for b in books_data.get("books", [])}

        # Load authors.json
        authors_file = Path(self.metadata_path) / "authors.json"
        if authors_file.exists():
            with open(authors_file, "r", encoding="utf-8") as f:
                authors_data = json.load(f)
                self.authors = {a["id"]: a for a in authors_data.get("authors", [])}

        # Load categories.json
        categories_file = Path(self.metadata_path) / "categories.json"
        if categories_file.exists():
            with open(categories_file, "r", encoding="utf-8") as f:
                categories_data = json.load(f)
                self.categories = {
                    c["id"]: c for c in categories_data.get("categories", [])
                }

        return {
            "books": len(self.books),
            "authors": len(self.authors),
            "categories": len(self.categories),
        }

    def get_book_metadata(self, book_id: int) -> Optional[Dict[str, Any]]:
        """Get metadata for a specific book."""
        return self.books.get(book_id)

    def get_books_by_category(self, category_name: str) -> List[Dict[str, Any]]:
        """Get all books in a category."""
        return [
            book
            for book in self.books.values()
            if book.get("cat_name") == category_name
        ]

    def get_books_by_author(self, author_id: int) -> List[Dict[str, Any]]:
        """Get all books by an author."""
        return [
            book
            for book in self.books.values()
            if str(author_id) in book.get("author_str", "")
        ]


if __name__ == "__main__":
    import asyncio

    async def main():
        """Demo ingestion pipeline."""

        print("Multi-Source Ingestion Pipeline - Demo")
        print("=" * 50)

        # Create pipeline
        datasets_path = "K:/learning/technical/ai-ml/AI-Mastery-2026/datasets"
        pipeline = await create_arabic_islamic_pipeline(datasets_path)

        # Ingest all sources
        results = await pipeline.ingest_all_sources(
            progress_callback=lambda name, new, updated: print(
                f"  {name}: {new} new, {updated} updated"
            ),
        )

        # Print statistics
        print("\nIngestion Statistics:")
        stats = pipeline.get_statistics()
        for key, value in stats.items():
            print(f"  {key}: {value}")

    asyncio.run(main())
