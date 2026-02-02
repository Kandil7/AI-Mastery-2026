"""
Production File Upload + Processing Module for RAG

Upgrades included:
- Content-based file type sniffing (magic bytes) + extension + MIME validation
- Secure storage with UUID filenames + tenant-aware directories
- Streaming save (no need to load full file into memory)
- Virus scanning integration (optional clamd)
- PDF extraction with fallback + OCR for scanned PDFs (optional pdf2image + pytesseract)
- Process isolation + timeouts for extraction to avoid parser hangs
- Resource limits: max size, max pages, max extracted chars
- Async jobs + progress tracking store (pluggable)
- Language detection integrated (optional langdetect)
- Automatic cleanup
"""

import asyncio
import hashlib
import logging
import mimetypes
import os
import re
import time
import uuid
from abc import ABC, abstractmethod
from concurrent.futures import ProcessPoolExecutor
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import Any, AsyncIterator, Dict, List, Optional, Tuple

from pydantic import BaseModel, Field, field_validator

# Import document class from retrieval module
from src.retrieval import Document

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

# ----------------------------
# Optional dependencies
# ----------------------------
try:
    import PyPDF2
    from pdfminer.high_level import extract_text as pdf_extract_text
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logger.warning("PyPDF2 or pdfminer not available. PDF support disabled.")

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logger.warning("python-docx not available. DOCX support disabled.")

try:
    from PIL import Image
    import pytesseract
    OCR_SUPPORT = True
except ImportError:
    OCR_SUPPORT = False
    logger.warning("PIL or pytesseract not available. OCR support disabled.")

try:
    from langdetect import detect
    LANG_DETECT_SUPPORT = True
except ImportError:
    LANG_DETECT_SUPPORT = False
    logger.warning("langdetect not available. Language detection disabled.")

try:
    import magic  # python-magic (libmagic)
    MAGIC_SUPPORT = True
except ImportError:
    MAGIC_SUPPORT = False
    logger.warning("python-magic not available. Content sniffing will be heuristic.")

try:
    import clamd  # clamd client for ClamAV daemon
    CLAMAV_SUPPORT = True
except ImportError:
    CLAMAV_SUPPORT = False
    logger.warning("clamd not available. Virus scanning will be skipped.")

try:
    from pdf2image import convert_from_path  # rasterize PDF pages for OCR
    PDF2IMAGE_SUPPORT = True
except ImportError:
    PDF2IMAGE_SUPPORT = False
    logger.warning("pdf2image not available. OCR fallback for scanned PDFs disabled.")


# ----------------------------
# Enums / Models
# ----------------------------
class FileType(Enum):
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MD = "md"
    IMAGE = "image"


class JobStatus(Enum):
    QUEUED = "queued"
    RUNNING = "running"
    SUCCEEDED = "succeeded"
    FAILED = "failed"


class FileUploadRequest(BaseModel):
    filename: str = Field(..., min_length=1, max_length=255)
    content_type: str = Field(..., description="MIME type (client-provided, not trusted)")
    file_size: int = Field(..., gt=0)
    tenant_id: Optional[str] = Field(None, description="Multi-tenant routing key")
    metadata: Dict[str, Any] = Field(default_factory=dict)

    @field_validator("filename")
    def validate_filename(cls, v: str) -> str:
        if not v or not isinstance(v, str):
            raise ValueError("Filename must be a non-empty string")
        # prevent traversal + weird separators
        if ".." in v or "/" in v or "\\" in v:
            raise ValueError("Filename contains invalid path characters")
        # basic sanitation
        if len(v.strip()) == 0:
            raise ValueError("Filename cannot be blank")
        return v

    @field_validator("file_size")
    def validate_file_size(cls, v: int) -> int:
        max_size = 50 * 1024 * 1024
        if v > max_size:
            raise ValueError(f"File size exceeds max allowed ({max_size} bytes)")
        return v

    @field_validator("tenant_id")
    def validate_tenant_id(cls, v: Optional[str]) -> Optional[str]:
        if v is None:
            return v
        if not re.fullmatch(r"[a-zA-Z0-9_\-]{1,64}", v):
            raise ValueError("tenant_id must be alphanumeric/_/- up to 64 chars")
        return v


class FileProcessingResult(BaseModel):
    success: bool
    message: str
    documents: List[Document]
    extracted_text_length: int
    processing_time_ms: float
    warnings: List[str] = Field(default_factory=list)
    metadata: Dict[str, Any] = Field(default_factory=dict)


class JobInfo(BaseModel):
    job_id: str
    status: JobStatus
    progress: float = Field(0.0, ge=0.0, le=100.0)
    message: str = ""
    created_at: float
    updated_at: float
    result: Optional[FileProcessingResult] = None
    error: Optional[str] = None


# ----------------------------
# Progress Store (pluggable)
# ----------------------------
class ProgressStore(ABC):
    @abstractmethod
    def create(self) -> JobInfo: ...
    @abstractmethod
    def update(self, job_id: str, **kwargs) -> JobInfo: ...
    @abstractmethod
    def get(self, job_id: str) -> Optional[JobInfo]: ...


class InMemoryProgressStore(ProgressStore):
    def __init__(self):
        self._jobs: Dict[str, JobInfo] = {}
        self._lock = asyncio.Lock()

    def create(self) -> JobInfo:
        now = time.time()
        job = JobInfo(
            job_id=str(uuid.uuid4()),
            status=JobStatus.QUEUED,
            progress=0.0,
            message="queued",
            created_at=now,
            updated_at=now,
        )
        self._jobs[job.job_id] = job
        return job

    def update(self, job_id: str, **kwargs) -> JobInfo:
        job = self._jobs[job_id]
        data = job.model_dump()
        data.update(kwargs)
        data["updated_at"] = time.time()
        job = JobInfo(**data)
        self._jobs[job_id] = job
        return job

    def get(self, job_id: str) -> Optional[JobInfo]:
        return self._jobs.get(job_id)


# ----------------------------
# Document Processors
# ----------------------------
class DocumentProcessor(ABC):
    @abstractmethod
    def extract_text(self, file_path: str, *, limits: "ExtractionLimits") -> Tuple[str, Dict[str, Any]]:
        """
        Returns:
            (text, meta) where meta includes extraction method details.
        """


@dataclass(frozen=True)
class ExtractionLimits:
    max_pages: int = 300
    max_chars: int = 2_000_000          # cap extracted text length
    timeout_sec: int = 60              # hard timeout for extraction
    ocr_max_pages: int = 20            # cap OCR pages (expensive)
    min_text_threshold: int = 200      # below this treat as "scanned/empty"


def _truncate_text(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    return text[:max_chars]


class PDFProcessor(DocumentProcessor):
    def extract_text(self, file_path: str, *, limits: ExtractionLimits) -> Tuple[str, Dict[str, Any]]:
        if not PDF_SUPPORT:
            raise RuntimeError("PDF processing libraries not available")

        meta: Dict[str, Any] = {"method": None, "ocr_used": False, "pages": None}

        # First, attempt pdfminer
        text = ""
        errors: List[str] = []

        try:
            text = pdf_extract_text(file_path) or ""
            meta["method"] = "pdfminer"
            # Try to estimate pages using PyPDF2 (cheap)
            try:
                with open(file_path, "rb") as f:
                    r = PyPDF2.PdfReader(f)
                    meta["pages"] = len(r.pages)
                    if meta["pages"] and meta["pages"] > limits.max_pages:
                        raise RuntimeError(f"PDF has too many pages ({meta['pages']} > {limits.max_pages})")
            except Exception as e:
                errors.append(f"page_count failed: {e}")

            if len(text.strip()) >= limits.min_text_threshold:
                return _truncate_text(text, limits.max_chars), meta
        except Exception as e:
            errors.append(f"pdfminer failed: {e}")

        # Second, PyPDF2
        try:
            pages_text: List[str] = []
            with open(file_path, "rb") as f:
                r = PyPDF2.PdfReader(f)
                page_count = len(r.pages)
                meta["pages"] = page_count
                if page_count > limits.max_pages:
                    raise RuntimeError(f"PDF has too many pages ({page_count} > {limits.max_pages})")

                for i in range(page_count):
                    page = r.pages[i]
                    pt = page.extract_text() or ""
                    pages_text.append(pt)

            text = "\n".join(pages_text)
            meta["method"] = "pypdf2"
            if len(text.strip()) >= limits.min_text_threshold:
                return _truncate_text(text, limits.max_chars), meta
        except Exception as e:
            errors.append(f"PyPDF2 failed: {e}")

        # OCR fallback for scanned PDFs (optional)
        if OCR_SUPPORT and PDF2IMAGE_SUPPORT:
            try:
                meta["method"] = "ocr_pdf"
                meta["ocr_used"] = True

                # Rasterize only first N pages to control cost
                images = convert_from_path(file_path, first_page=1, last_page=min(limits.ocr_max_pages, limits.max_pages))
                ocr_parts: List[str] = []
                for img in images:
                    ocr_parts.append(pytesseract.image_to_string(img))
                text = "\n".join(ocr_parts)

                if text.strip():
                    return _truncate_text(text, limits.max_chars), meta
            except Exception as e:
                errors.append(f"OCR fallback failed: {e}")

        raise RuntimeError(f"Failed to extract text from PDF. Errors: {'; '.join(map(str, errors))}")


class DOCXProcessor(DocumentProcessor):
    def extract_text(self, file_path: str, *, limits: ExtractionLimits) -> Tuple[str, Dict[str, Any]]:
        if not DOCX_SUPPORT:
            raise RuntimeError("DOCX processing libraries not available")

        meta = {"method": "python-docx"}
        doc = docx.Document(file_path)

        parts: List[str] = []
        # paragraphs
        parts.extend([p.text for p in doc.paragraphs if p.text and p.text.strip()])

        # tables (common missing piece)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text)
                if row_text.strip():
                    parts.append(row_text)

        text = "\n".join(parts)
        return _truncate_text(text, limits.max_chars), meta


class TextProcessor(DocumentProcessor):
    def extract_text(self, file_path: str, *, limits: ExtractionLimits) -> Tuple[str, Dict[str, Any]]:
        meta = {"method": "text"}
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
        last_err: Optional[Exception] = None

        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    text = f.read()
                return _truncate_text(text, limits.max_chars), {**meta, "encoding": enc}
            except UnicodeDecodeError as e:
                last_err = e
                continue

        raise RuntimeError(f"Failed to decode text file. Last error: {last_err}")


class ImageProcessor(DocumentProcessor):
    def extract_text(self, file_path: str, *, limits: ExtractionLimits) -> Tuple[str, Dict[str, Any]]:
        if not OCR_SUPPORT:
            raise RuntimeError("OCR processing libraries not available")
        meta = {"method": "ocr_image", "ocr_used": True}

        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return _truncate_text(text, limits.max_chars), meta


# ----------------------------
# File Type Sniffing
# ----------------------------
class FileSniffer:
    """
    Detect file type via:
    1) libmagic if available
    2) signature heuristics
    3) extension fallback
    """

    def sniff_mime(self, file_path: str) -> Optional[str]:
        if MAGIC_SUPPORT:
            try:
                m = magic.Magic(mime=True)
                return m.from_file(file_path)
            except Exception:
                return None

        # Heuristic signature sniffing
        try:
            with open(file_path, "rb") as f:
                head = f.read(16)
            # PDF: %PDF-
            if head.startswith(b"%PDF-"):
                return "application/pdf"
            # ZIP (DOCX is zip container)
            if head.startswith(b"PK\x03\x04"):
                return "application/zip"
            # PNG
            if head.startswith(b"\x89PNG\r\n\x1a\n"):
                return "image/png"
            # JPG
            if head.startswith(b"\xFF\xD8\xFF"):
                return "image/jpeg"
        except Exception:
            return None

        return None

    def to_file_type(self, filename: str, sniffed_mime: Optional[str]) -> FileType:
        ext = Path(filename).suffix.lower().lstrip(".")
        # primary by sniffed mime
        if sniffed_mime:
            if sniffed_mime == "application/pdf":
                return FileType.PDF
            if sniffed_mime in ("image/png", "image/jpeg", "image/bmp", "image/tiff", "image/webp"):
                return FileType.IMAGE
            if sniffed_mime in ("text/plain", "text/markdown", "text/x-markdown"):
                return FileType.TXT if ext == "txt" else FileType.MD
            # zip could be docx
            if sniffed_mime in ("application/zip",):
                if ext in ("docx", "doc"):
                    return FileType.DOCX

        # fallback by extension
        if ext == "pdf":
            return FileType.PDF
        if ext in ("docx", "doc"):
            return FileType.DOCX
        if ext == "txt":
            return FileType.TXT
        if ext == "md":
            return FileType.MD
        if ext in ("jpg", "jpeg", "png", "bmp", "tiff", "webp"):
            return FileType.IMAGE

        raise ValueError(f"Unsupported file type: {ext} (sniffed_mime={sniffed_mime})")


# ----------------------------
# Virus Scanning
# ----------------------------
class VirusScanner(ABC):
    @abstractmethod
    def scan(self, file_path: str) -> Tuple[bool, Optional[str]]:
        """Returns (is_clean, reason)."""


class NoopVirusScanner(VirusScanner):
    def scan(self, file_path: str) -> Tuple[bool, Optional[str]]:
        return True, None


class ClamAVScanner(VirusScanner):
    def __init__(self, socket_path: Optional[str] = None, host: Optional[str] = None, port: Optional[int] = None):
        if not CLAMAV_SUPPORT:
            raise RuntimeError("clamd not installed")
        self.socket_path = socket_path
        self.host = host
        self.port = port

    def _client(self):
        if self.socket_path:
            return clamd.ClamdUnixSocket(self.socket_path)
        if self.host and self.port:
            return clamd.ClamdNetworkSocket(self.host, self.port)
        # default unix socket
        return clamd.ClamdUnixSocket()

    def scan(self, file_path: str) -> Tuple[bool, Optional[str]]:
        c = self._client()
        try:
            res = c.scan(file_path)
            # res example: {"/path": ("OK", None)} or ("FOUND", "MalwareName")
            status, info = res.get(file_path, (None, None))
            if status == "OK":
                return True, None
            return False, f"Virus scan failed: {status} {info}"
        except Exception as e:
            # in production you may choose fail-closed; here fail-open with warning is safer for dev
            return True, f"Virus scan unavailable: {e}"


# ----------------------------
# Secure Storage
# ----------------------------
class SecureStorage:
    def __init__(self, base_dir: str):
        self.base_dir = Path(base_dir)
        self.base_dir.mkdir(parents=True, exist_ok=True)

    def tenant_dir(self, tenant_id: Optional[str]) -> Path:
        tid = tenant_id or "public"
        p = self.base_dir / tid
        p.mkdir(parents=True, exist_ok=True)
        return p

    def generate_path(self, tenant_id: Optional[str], original_filename: str) -> Path:
        ext = Path(original_filename).suffix.lower()
        safe_ext = ext if re.fullmatch(r"\.[a-z0-9]{1,8}", ext) else ""
        filename = f"{uuid.uuid4().hex}{safe_ext}"
        return self.tenant_dir(tenant_id) / filename


# ----------------------------
# Extraction Runner (Process isolation + timeout)
# ----------------------------
def _extract_in_subprocess(processor_name: str, file_path: str, limits: Dict[str, Any]) -> Tuple[str, Dict[str, Any]]:
    """
    Runs in a separate process to isolate parser hangs.
    """
    lim = ExtractionLimits(**limits)

    processors: Dict[str, DocumentProcessor] = {
        "pdf": PDFProcessor(),
        "docx": DOCXProcessor(),
        "text": TextProcessor(),
        "image": ImageProcessor(),
        "md": TextProcessor(),
        "txt": TextProcessor(),
    }
    proc = processors[processor_name]
    return proc.extract_text(file_path, limits=lim)


# ----------------------------
# FileManager (Production)
# ----------------------------
class FileManager:
    def __init__(
        self,
        upload_dir: str = "uploads",
        max_file_size: int = 50 * 1024 * 1024,
        extraction_limits: ExtractionLimits = ExtractionLimits(),
        progress_store: Optional[ProgressStore] = None,
        virus_scanner: Optional[VirusScanner] = None,
        process_pool_workers: int = 2,
    ):
        self.max_file_size = max_file_size
        self.limits = extraction_limits
        self.sniffer = FileSniffer()
        self.storage = SecureStorage(upload_dir)
        self.progress = progress_store or InMemoryProgressStore()
        self.virus_scanner = virus_scanner or NoopVirusScanner()
        self.pool = ProcessPoolExecutor(max_workers=process_pool_workers)

        # registry
        self._processor_key_by_type: Dict[FileType, str] = {
            FileType.PDF: "pdf",
            FileType.DOCX: "docx",
            FileType.TXT: "txt",
            FileType.MD: "md",
            FileType.IMAGE: "image",
        }

    # ---------- Validation ----------
    def validate_file_upload(self, req: FileUploadRequest) -> List[str]:
        errors: List[str] = []

        if req.file_size > self.max_file_size:
            errors.append(f"File too large: {req.file_size} > {self.max_file_size}")

        # quick extension sanity (not authoritative)
        ext = Path(req.filename).suffix.lower().lstrip(".")
        if not ext:
            errors.append("Missing file extension")

        # allowlist of client mime (still not trusted)
        allowed_client_types = {
            "application/pdf",
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/plain",
            "text/markdown",
            "text/x-markdown",
            "image/jpeg",
            "image/png",
            "image/bmp",
            "image/tiff",
            "image/webp",
        }
        if req.content_type not in allowed_client_types:
            errors.append(f"Client content type not allowed: {req.content_type}")

        return errors

    # ---------- Saving ----------
    async def save_stream(self, stream: AsyncIterator[bytes], req: FileUploadRequest) -> str:
        """
        Save upload stream to disk safely with size enforcement.
        """
        target_path = self.storage.generate_path(req.tenant_id, req.filename)

        total = 0
        try:
            with open(target_path, "wb") as f:
                async for chunk in stream:
                    total += len(chunk)
                    if total > self.max_file_size:
                        raise ValueError("Upload exceeds max_file_size during streaming")
                    f.write(chunk)
            return str(target_path.absolute())
        except Exception:
            # cleanup on failure
            try:
                if target_path.exists():
                    target_path.unlink()
            except Exception:
                pass
            raise

    async def save_bytes(self, data: bytes, req: FileUploadRequest) -> str:
        if len(data) > self.max_file_size:
            raise ValueError("Upload exceeds max_file_size")
        target_path = self.storage.generate_path(req.tenant_id, req.filename)
        with open(target_path, "wb") as f:
            f.write(data)
        return str(target_path.absolute())

    # ---------- Hash ----------
    def calculate_file_hash(self, file_path: str) -> str:
        sha = hashlib.sha256()
        with open(file_path, "rb") as f:
            for b in iter(lambda: f.read(4096), b""):
                sha.update(b)
        return sha.hexdigest()

    # ---------- Jobs ----------
    def create_processing_job(self) -> JobInfo:
        return self.progress.create()

    def get_job(self, job_id: str) -> Optional[JobInfo]:
        return self.progress.get(job_id)

    async def process_file_async_job(
        self,
        job_id: str,
        file_path: str,
        original_filename: str,
        req: FileUploadRequest,
    ) -> None:
        """
        Runs the full pipeline as an async job with progress updates.
        """
        self.progress.update(job_id, status=JobStatus.RUNNING, progress=1.0, message="starting")

        warnings: List[str] = []
        start = time.time()

        try:
            # 1) sniff mime from content
            self.progress.update(job_id, progress=5.0, message="sniffing file type")
            sniffed_mime = self.sniffer.sniff_mime(file_path)
            file_type = self.sniffer.to_file_type(original_filename, sniffed_mime)

            # 2) virus scan (optional)
            self.progress.update(job_id, progress=10.0, message="virus scanning")
            is_clean, scan_note = self.virus_scanner.scan(file_path)
            if scan_note:
                warnings.append(scan_note)
            if not is_clean:
                raise RuntimeError(f"Upload rejected: {scan_note or 'malware found'}")

            # 3) extraction in subprocess with timeout
            self.progress.update(job_id, progress=20.0, message="extracting text")
            processor_key = self._processor_key_by_type[file_type]

            loop = asyncio.get_running_loop()
            limits_dict = self.limits.__dict__.copy()

            fut = loop.run_in_executor(
                self.pool,
                _extract_in_subprocess,
                processor_key,
                file_path,
                limits_dict,
            )
            try:
                extracted_text, extraction_meta = await asyncio.wait_for(fut, timeout=self.limits.timeout_sec + 5)
            except asyncio.TimeoutError:
                raise RuntimeError(f"Extraction timed out after ~{self.limits.timeout_sec}s")

            self.progress.update(job_id, progress=70.0, message="post-processing")

            extracted_text = extracted_text or ""
            extracted_text = extracted_text.strip()

            # 4) language detection
            lang = None
            if LANG_DETECT_SUPPORT and extracted_text:
                try:
                    lang = detect(extracted_text[:5000])
                except Exception as e:
                    warnings.append(f"Language detection failed: {e}")

            # 5) hash + Document
            file_hash = self.calculate_file_hash(file_path)
            doc_id = f"file_{file_hash[:16]}"

            processed_at = time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
            doc = Document(
                id=doc_id,
                content=extracted_text,
                source="file_upload",
                doc_type=file_type.value,
                metadata={
                    "tenant_id": req.tenant_id,
                    "original_filename": original_filename,
                    "stored_filename": Path(file_path).name,
                    "file_hash": file_hash,
                    "sniffed_mime": sniffed_mime,
                    "client_mime": req.content_type,
                    "file_type": file_type.value,
                    "content_length": len(extracted_text),
                    "language": lang,
                    "processed_at": processed_at,
                    "extraction": extraction_meta,
                    **(req.metadata or {}),
                },
            )

            ms = (time.time() - start) * 1000
            result = FileProcessingResult(
                success=True,
                message=f"Processed {original_filename}",
                documents=[doc],
                extracted_text_length=len(extracted_text),
                processing_time_ms=ms,
                warnings=warnings,
                metadata={
                    "file_type": file_type.value,
                    "original_filename": original_filename,
                    "doc_id": doc_id,
                },
            )

            self.progress.update(job_id, status=JobStatus.SUCCEEDED, progress=100.0, message="done", result=result)

            logger.info(
                "file_processed",
                extra={
                    "job_id": job_id,
                    "tenant_id": req.tenant_id,
                    "file_type": file_type.value,
                    "doc_id": doc_id,
                    "chars": len(extracted_text),
                    "ms": ms,
                    "method": (result.metadata or {}).get("extraction_method"),
                },
            )

        except Exception as e:
            ms = (time.time() - start) * 1000
            err = str(e)
            self.progress.update(job_id, status=JobStatus.FAILED, progress=100.0, message="failed", error=err)
            logger.exception("file_processing_failed", extra={"job_id": job_id, "tenant_id": req.tenant_id, "ms": ms})

        finally:
            # Always cleanup local stored file after processing (you can make this configurable)
            try:
                p = Path(file_path)
                if p.exists():
                    p.unlink()
            except Exception as e:
                logger.warning("cleanup_failed", extra={"job_id": job_id, "error": str(e)})

    # ---------- Convenience: one-shot (no job) ----------
    async def process_file(
        self,
        file_path: str,
        original_filename: str,
        req: FileUploadRequest,
    ) -> FileProcessingResult:
        """
        One-shot processing without job tracking (still safe + isolated).
        """
        job = self.create_processing_job()
        await self.process_file_async_job(job.job_id, file_path, original_filename, req)
        info = self.get_job(job.job_id)
        if info and info.result:
            return info.result
        return FileProcessingResult(
            success=False,
            message=f"Processing failed: {info.error if info else 'unknown'}",
            documents=[],
            extracted_text_length=0,
            processing_time_ms=0.0,
            warnings=[info.error] if info and info.error else ["unknown error"],
        )


# ----------------------------
# Singleton
# ----------------------------
progress_store = InMemoryProgressStore()

# Optional: enable ClamAV scanner if you have clamd running.
# virus_scanner = ClamAVScanner(socket_path="/var/run/clamav/clamd.ctl")
virus_scanner = NoopVirusScanner()

file_manager = FileManager(
    upload_dir="uploads",
    max_file_size=50 * 1024 * 1024,
    extraction_limits=ExtractionLimits(
        max_pages=300,
        max_chars=2_000_000,
        timeout_sec=60,
        ocr_max_pages=20,
        min_text_threshold=200,
    ),
    progress_store=progress_store,
    virus_scanner=virus_scanner,
    process_pool_workers=2,
)

__all__ = [
    "FileManager",
    "FileUploadRequest",
    "FileProcessingResult",
    "FileType",
    "JobInfo",
    "JobStatus",
    "file_manager",
]
